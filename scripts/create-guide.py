#!/usr/bin/env python3
"""Generate CGA Project Guide DOCX document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Color Palette (Cool + Active, Tech/Modern) ──────────────────────────
PRIMARY = RGBColor(0x0F, 0x17, 0x2A)       # Deep navy — headings
BODY_CLR = RGBColor(0x1E, 0x29, 0x3B)      # Dark slate — body text
SECONDARY = RGBColor(0x64, 0x74, 0x8B)     # Muted gray — captions
ACCENT = RGBColor(0xEF, 0x44, 0x44)        # Red-500 — accent (CGA brand)
ACCENT_DARK = RGBColor(0xDC, 0x26, 0x26)   # Red-600 — links/strong accent
TABLE_HEADER_BG = "0F172A"                 # Dark navy
TABLE_ALT_BG = "F8FAFC"                    # Very light blue
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

FONT = "Calibri"
LINE_SPACING = 1.3

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, inside_h=None, inside_v=None):
    """Set cell borders individually."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val["style"]}" w:sz="{val["size"]}" w:color="{val["color"]}" w:space="0"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(paragraph, text, bold=False, italic=False, size=11, color=BODY_CLR, font=FONT, underline=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    if underline:
        run.font.underline = True
    # Set East Asian font too
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), font)
    return run


def set_line_spacing(paragraph, spacing=LINE_SPACING):
    """Set line spacing multiplier on a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    spacing_elem = pPr.find(qn("w:spacing"))
    if spacing_elem is None:
        spacing_elem = parse_xml(f'<w:spacing {nsdecls("w")} w:line="{int(spacing * 240)}" w:lineRule="auto"/>')
        pPr.append(spacing_elem)
    else:
        spacing_elem.set(qn("w:line"), str(int(spacing * 240)))
        spacing_elem.set(qn("w:lineRule"), "auto")


def add_heading(doc, text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(level=level)
    h.clear()
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: PRIMARY, 2: PRIMARY, 3: BODY_CLR}
    add_run(h, text, bold=True, size=sizes.get(level, 12), color=colors.get(level, PRIMARY))
    set_line_spacing(h, LINE_SPACING)
    # Remove spacing after
    pPr = h._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="160" w:line="{int(LINE_SPACING * 240)}" w:lineRule="auto"/>')
        pPr.append(spacing)
    else:
        spacing.set(qn("w:after"), "160")
    return h


def add_body(doc, text, indent=True, bold=False, italic=False, size=11, color=BODY_CLR):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    set_line_spacing(p, LINE_SPACING)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code_block(doc, code_text):
    """Add a code-like block (indented, monospace-styled)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Consolas" w:hAnsi="Consolas" w:eastAsia="Consolas"/>')
    # Replace existing rFonts
    old = rPr.find(qn("w:rFonts"))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, rFonts)
    set_line_spacing(p, 1.15)
    return p


def add_step(doc, number, text):
    """Add a numbered step paragraph."""
    p = doc.add_paragraph()
    add_run(p, f"{number}. ", bold=True, size=11, color=ACCENT)
    add_run(p, text, size=11, color=BODY_CLR)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(0)
    set_line_spacing(p, LINE_SPACING)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11, color=BODY_CLR)
        add_run(p, text, size=11, color=BODY_CLR)
    else:
        add_run(p, text, size=11, color=BODY_CLR)
    set_line_spacing(p, LINE_SPACING)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_run(p, header_text, bold=True, size=10, color=WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_line_spacing(p, LINE_SPACING)
        set_cell_shading(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Style data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_run(p, str(cell_text), size=10, color=BODY_CLR)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(p, LINE_SPACING)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)

    # Set table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="E2E8F0"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    # Remove existing borders
    old_borders = tblPr.find(qn("w:tblBorders"))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblPr.append(borders)

    # Set cell margins
    tblCellMar = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="60" w:type="dxa"/>'
        f'  <w:left w:w="120" w:type="dxa"/>'
        f'  <w:bottom w:w="60" w:type="dxa"/>'
        f'  <w:right w:w="120" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    old_mar = tblPr.find(qn("w:tblCellMar"))
    if old_mar is not None:
        tblPr.remove(old_mar)
    tblPr.append(tblCellMar)

    return table


def add_empty_line(doc):
    """Add a small spacer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("")
    run.font.size = Pt(6)
    return p


# ── Document Setup ──────────────────────────────────────────────────────

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# Set default font
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.font.color.rgb = BODY_CLR
style.paragraph_format.line_spacing = LINE_SPACING
rPr_default = style.element.get_or_add_rPr()
rFonts_default = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}" w:eastAsia="{FONT}"/>')
old_rf = rPr_default.find(qn("w:rFonts"))
if old_rf is not None:
    rPr_default.remove(old_rf)
rPr_default.insert(0, rFonts_default)

# Configure heading styles
for level, (sz, clr, space_before, space_after) in {
    1: (18, PRIMARY, 360, 160),
    2: (14, PRIMARY, 280, 120),
    3: (12, BODY_CLR, 200, 100),
}.items():
    h_style = doc.styles[f"Heading {level}"]
    h_style.font.name = FONT
    h_style.font.size = Pt(sz)
    h_style.font.color.rgb = clr
    h_style.font.bold = True
    h_style.paragraph_format.line_spacing = LINE_SPACING
    h_style.paragraph_format.space_before = Pt(space_before / 20)
    h_style.paragraph_format.space_after = Pt(space_after / 20)
    h_rPr = h_style.element.get_or_add_rPr()
    h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}" w:eastAsia="{FONT}"/>')
    old_hrf = h_rPr.find(qn("w:rFonts"))
    if old_hrf is not None:
        h_rPr.remove(old_hrf)
    h_rPr.insert(0, h_rFonts)

# ── Title ────────────────────────────────────────────────────────────────

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title_p, "Create Growth Agency", bold=True, size=26, color=PRIMARY)
set_line_spacing(title_p, LINE_SPACING)
title_p.paragraph_format.space_after = Pt(0)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(subtitle_p, "Complete Project Guide", bold=True, size=16, color=ACCENT)
set_line_spacing(subtitle_p, LINE_SPACING)
subtitle_p.paragraph_format.space_after = Pt(4)

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(desc_p, "How to Manage, Customize & Go Live", italic=True, size=12, color=SECONDARY)
set_line_spacing(desc_p, LINE_SPACING)
desc_p.paragraph_format.space_after = Pt(200)

# ── Section 1: Project Overview ─────────────────────────────────────────

add_heading(doc, "1. Project Overview", level=1)

add_body(doc, (
    "Create Growth Agency (CGA) is a full-featured YouTube content creation SaaS platform designed "
    "specifically for Bangladeshi content creators. The platform provides a comprehensive suite of "
    "services including YouTube video editing, thumbnail design, channel management, SEO optimization, "
    "scriptwriting, and content strategy planning. Clients can browse available services, place orders, "
    "make payments via SSLCOMMERZ (the leading Bangladeshi payment gateway), and track their order "
    "status through a personalized dashboard."
))

add_body(doc, (
    "The project is built on a modern, production-ready tech stack. The frontend and backend are powered "
    "by Next.js 16 with TypeScript, providing type safety, server-side rendering, and API route "
    "capabilities in a single unified framework. Data persistence is handled through Prisma ORM connected "
    "to a Supabase PostgreSQL database, offering robust relational data management with real-time "
    "capabilities. The user interface is crafted with Tailwind CSS and the shadcn/ui component library, "
    "delivering a polished, responsive design that works seamlessly across all device sizes. Framer Motion "
    "adds smooth animations and transitions throughout the application, enhancing the overall user "
    "experience."
))

add_body(doc, (
    "As of the current development milestone, Days 1 through 5 are fully complete. This includes the "
    "complete authentication system (registration, login, session management with JWT), the service "
    "catalog with dynamic pages for each offering, a fully functional order management system with "
    "payment integration through SSLCOMMERZ, and a comprehensive admin panel for managing users, "
    "orders, services, and platform settings. The remaining work encompasses Days 6 through 14, which "
    "will deliver a blog and content management system, real-time chat support, an AI-powered content "
    "assistant widget, Progressive Web App (PWA) capabilities, internationalization support for both "
    "Bengali and English, comprehensive testing and quality assurance, and the final production deployment."
))

# ── Section 2: Current File Structure ───────────────────────────────────

add_heading(doc, "2. Current File Structure", level=1)

add_body(doc, (
    "The project follows a well-organized directory structure common to Next.js applications. "
    "Below is a summary of the key directories and their purposes."
), indent=False)

structure_headers = ["Directory / File", "Purpose"]
structure_rows = [
    ["src/app/", "All pages and API routes (Next.js App Router)"],
    ["src/components/ui/", "Reusable UI primitives from shadcn/ui"],
    ["src/components/dashboard/", "Dashboard-specific components (sidebar, cards, etc.)"],
    ["src/components/admin/", "Admin panel components (admin-sidebar, tables, forms)"],
    ["src/components/services/", "Service display and selection components"],
    ["src/components/auth/", "Login, registration, and auth-related components"],
    ["src/components/layout/", "Shared layout components (navbar, footer)"],
    ["src/lib/", "Core utilities (db.ts, auth.ts, order-utils.ts, supabase.ts)"],
    ["src/hooks/", "Custom React hooks (e.g., use-auth.ts)"],
    ["prisma/schema.prisma", "Database schema definition (all models and relations)"],
    [".env", "Environment variables (NEVER commit this file to Git)"],
]
add_table(doc, structure_headers, structure_rows, col_widths=[2.5, 4.5])
add_empty_line(doc)

# ── Section 3: Environment Variables Guide ──────────────────────────────

add_heading(doc, "3. Environment Variables Guide", level=1)

add_body(doc, (
    "The application requires several environment variables to function correctly. These configure "
    "database connections, authentication, and third-party service integrations. All variables must be "
    "set in your local .env file during development and in the Vercel environment settings for "
    "production deployment."
), indent=False)

env_headers = ["Variable", "Purpose", "Where to Get It"]
env_rows = [
    ["DATABASE_URL",
     "Supabase PostgreSQL connection string",
     "Supabase Dashboard > Settings > Database > Connection string"],
    ["NEXT_PUBLIC_SUPABASE_URL",
     "Supabase project URL",
     "Supabase Dashboard > Settings > API"],
    ["NEXT_PUBLIC_SUPABASE_ANON_KEY",
     "Public API key (safe for client-side)",
     "Supabase Dashboard > Settings > API"],
    ["SUPABASE_SERVICE_ROLE_KEY",
     "Admin API key (NEVER expose to client)",
     "Supabase Dashboard > Settings > API"],
    ["JWT_SECRET",
     "JWT token signing secret",
     'Generate any random string: openssl rand -base64 32'],
    ["NEXT_PUBLIC_APP_URL",
     "Your live domain URL",
     "After deployment (e.g., https://yourdomain.com)"],
]
add_table(doc, env_headers, env_rows, col_widths=[2.2, 2.2, 2.6])
add_empty_line(doc)

# ── Section 4: How to Access Admin Panel ────────────────────────────────

add_heading(doc, "4. How to Access Admin Panel", level=1)

add_body(doc, (
    "The admin panel is the central control center for managing your CGA platform. Access is restricted "
    "to users with the admin role. Follow these steps to log in and navigate the admin area."
), indent=False)

add_step(doc, 1, "Open your web browser and navigate to your website URL (e.g., http://localhost:3000 during development or your live domain).")
add_step(doc, 2, "Log in with your admin credentials. Default credentials are: Email: admin@cga.com / Password: admin123. Change these immediately after first login.")
add_step(doc, 3, "After successful login, navigate to /admin in the URL bar. The admin panel will only load if your account has the role set to \"admin\" in the database.")
add_step(doc, 4, "Use the admin sidebar to navigate between the available admin pages: Dashboard (/admin) for an overview of key metrics, Users (/admin/users) for managing user accounts and roles, Orders (/admin/orders) for viewing and updating order statuses, Services (/admin/services) for editing service details and pricing, and Settings (/admin/settings) for platform-wide configuration.")
add_step(doc, 5, "Important: Only users whose role field is set to \"admin\" in the User table can access /admin routes. All other users will be redirected to their dashboard or the homepage.")

# ── Section 5: How to Change Text, Colors & Headlines ──────────────────

add_heading(doc, "5. How to Change Text, Colors & Headlines", level=1)

add_body(doc, (
    "One of the most common tasks after receiving the project is customizing the text, branding, and "
    "visual appearance to match your agency identity. This section provides a practical, file-by-file "
    "guide for making these changes. Each item below tells you exactly which file to open and what to "
    "modify."
))

add_heading(doc, "Site Name and Branding", level=2)
add_body(doc, (
    "The site title and meta description that appear in browser tabs and search engine results are "
    "defined in src/app/layout.tsx. Look for the metadata export object near the top of the file. "
    "Change the title field to your agency name and update the description field with your desired "
    "meta description. This affects SEO and browser tab display."
))

add_heading(doc, "Colors (Primary Accent)", level=2)
add_body(doc, (
    "The current primary accent color is red-500 (a vibrant red). To change this, you have two options. "
    "First, for a global replacement, you can update the Tailwind CSS configuration in tailwind.config.ts "
    "by extending the theme with your custom colors. Second, for CSS variable-based theming, modify the "
    "color definitions in src/app/globals.css. To find all instances of the current red-500 color, "
    "search your codebase for \"red-500\" or \"#EF4444\" and replace each occurrence with your preferred "
    "color. Common locations include buttons, links, badges, and decorative elements across multiple "
    "component files."
))

add_heading(doc, "Homepage Hero Text", level=2)
add_body(doc, (
    "The main heading and paragraph on the homepage are in src/app/page.tsx. Open this file and locate "
    "the hero section near the top. You will see large heading text and a descriptive paragraph below "
    "it. Simply edit the string content of these elements to update your homepage messaging."
))

add_heading(doc, "Service Descriptions", level=2)
add_body(doc, (
    "Service data (names, descriptions, prices, features) is stored in the database and seeded via "
    "the script at scripts/seed-services.ts. To change services permanently at the source, edit that "
    "seed file and re-run it. Alternatively, you can use the Admin Panel: navigate to Admin > Services, "
    "click Edit on any service, and modify the name, description, price, features list, and package "
    "details directly through the web interface. Changes made via the admin panel take effect immediately."
))

add_heading(doc, "Pricing Page", level=2)
add_body(doc, (
    "The pricing page at src/app/pricing/page.tsx contains hardcoded price values and feature lists "
    "displayed in pricing cards. Open this file and locate the pricing card data arrays or objects. "
    "Update the amounts (in BDT), feature descriptions, and any comparison details to reflect your "
    "actual pricing structure."
))

add_heading(doc, "Footer Text", level=2)
add_body(doc, (
    "The footer content is defined in src/components/layout/footer.tsx. This file contains your "
    "agency description, navigation links, social media links, and copyright notice. Edit the text "
    "directly in this component to update what appears at the bottom of every page."
))

add_heading(doc, "Navbar Links", level=2)
add_body(doc, (
    "Navigation bar items are configured in src/components/layout/navbar.tsx. You will find an array "
    "or list of navigation items with labels and href paths. Add, remove, or reorder items here to "
    "customize the main navigation menu."
))

add_heading(doc, "Dark Background Colors", level=2)
add_body(doc, (
    "The dark backgrounds used throughout the site use hex color codes like bg-[#080E1A] and "
    "bg-[#0B1120]. To change these, do a project-wide search for these hex values and replace them "
    "with your preferred dark background colors. You can also define these as CSS custom properties "
    "in globals.css for easier management."
))

add_heading(doc, "Fonts", level=2)
add_body(doc, (
    "The project uses Google Fonts loaded via next/font/google in src/app/layout.tsx. Currently, "
    "Outfit is used for English text and Hind Siliguri for Bengali text. To switch fonts, import a "
    "different font from next/font/google and replace the variable assignments in the layout file. "
    "Then update the className references on the root element."
))

# ── Section 6: How to Add Your Logo ────────────────────────────────────

add_heading(doc, "6. How to Add Your Logo", level=1)

add_body(doc, (
    "Follow these steps to replace the default CGA text logo with your own custom logo image across "
    "the entire application."
), indent=False)

add_step(doc, 1, "Prepare your logo file. Use a PNG format with a transparent background. The recommended size is 200x200 pixels. Ensure the logo looks clean at small sizes (40x40px) since it appears in navbars and sidebars.")
add_step(doc, 2, "Place your logo file in the public/ directory at the project root. Name it logo.png (i.e., public/logo.png). This makes it accessible at /logo.png from any page.")
add_step(doc, 3, "Open src/components/layout/navbar.tsx and locate the CGA text logo (usually a div or span element containing \"CGA\"). Replace it with the following Image component:")
add_code_block(doc, '<Image src="/logo.png" alt="CGA" width={40} height={40} className="rounded-xl" />')
add_step(doc, 4, "Add the Image import at the top of the file if not already present:")
add_code_block(doc, 'import Image from "next/image";')
add_step(doc, 5, "Repeat steps 3 and 4 for the following files to ensure the logo appears everywhere:")
add_bullet(doc, "src/components/dashboard/sidebar.tsx \u2014 Dashboard sidebar logo")
add_bullet(doc, "src/components/admin/admin-sidebar.tsx \u2014 Admin sidebar logo")
add_bullet(doc, "src/app/(auth)/login/login-form.tsx \u2014 Login page logo")
add_bullet(doc, "src/app/(auth)/register/page.tsx \u2014 Registration page logo")

# ── Section 7: How to Add Animations & Background Videos ───────────────

add_heading(doc, "7. How to Add Animations & Background Videos", level=1)

add_heading(doc, "Framer Motion Animations", level=2)
add_body(doc, (
    "Framer Motion is already installed and used throughout the project. Every page currently uses "
    "a fade-in + slide-up animation pattern: initial={{ opacity: 0, y: 20 }} and animate={{ opacity: 1, "
    "y: 0 }}. To add more animations, import the motion component from \"framer-motion\" and wrap any "
    "element you want to animate. You can add staggered animations for lists, hover effects on cards, "
    "scroll-triggered animations using the whileInView prop, and page transition effects using AnimatePresence."
))

add_heading(doc, "Animated Background Video", level=2)
add_body(doc, (
    "To add a video background behind the hero section, place your MP4 file in public/hero-bg.mp4. "
    "Then, in src/app/page.tsx, add a video element inside the hero section container (which should "
    "have position: relative). Insert the following code:"
))
add_code_block(doc, (
    '<video autoPlay muted loop playsInline\n'
    '  className="absolute inset-0 w-full h-full object-cover opacity-20">\n'
    '  <source src="/hero-bg.mp4" type="video/mp4" />\n'
    '</video>'
))
add_body(doc, (
    "The opacity-20 class ensures the video is subtle and does not interfere with text readability. "
    "Adjust the opacity value (0-100) to your preference. The autoPlay, muted, and loop attributes "
    "ensure the video plays automatically and continuously without sound."
))

add_heading(doc, "Lottie Animations", level=2)
add_body(doc, (
    "For lightweight vector animations, you can use Lottie. Install the package with npm install "
    "lottie-react, then import and use the Lottie component with a JSON animation file. Lottie "
    "animations are ideal for loading states, empty states, and decorative elements."
))

add_heading(doc, "CSS Animations", level=2)
add_body(doc, (
    "For simple keyframe animations (such as pulsing, floating, or gradient shifting), add custom "
    "@keyframes rules in src/app/globals.css. You can then reference these animations using Tailwind's "
    "animate utility or custom className properties."
))

# ── Section 8: How to Deploy to Live (Vercel) ──────────────────────────

add_heading(doc, "8. How to Deploy to Live (Vercel \u2014 Recommended)", level=1)

add_body(doc, (
    "Vercel is the recommended hosting platform for Next.js applications. It provides automatic "
    " deployments from Git, serverless functions, edge caching, and a generous free tier. Follow "
    "these steps to deploy your CGA project."
), indent=False)

add_step(doc, 1, "Push your code to a GitHub repository if you have not already. Ensure your .env file is listed in .gitignore and is NOT committed.")
add_step(doc, 2, "Go to vercel.com and sign in with your GitHub account. Click \"Add New\" > \"Project\" and select your CGA repository from the list.")
add_step(doc, 3, "In the Environment Variables section of the Vercel project settings, add ALL variables from your .env file. Use the exact same variable names. For JWT_SECRET, generate a NEW random string for production (do not reuse the development secret).")
add_step(doc, 4, "Set the Build Command to: npx prisma generate && next build. This ensures Prisma client is regenerated before the build. The Output Directory should be left as default (.next).")
add_step(doc, 5, "Click \"Deploy\". Vercel will build and deploy your application. You will receive a temporary .vercel.app URL. Test thoroughly before setting up a custom domain.")
add_step(doc, 6, "To add a custom domain, go to your Vercel project > Settings > Domains. Click \"Add Domain\" and enter your domain name (e.g., yourdomain.com). Follow the DNS configuration instructions provided.")
add_step(doc, 7, "Configure your domain DNS: add an A record pointing to 76.76.21.21, or a CNAME record pointing to cname.vercel-dns.com. DNS propagation typically takes a few minutes to a few hours.")

# ── Section 9: How to Manage After Going Live ──────────────────────────

add_heading(doc, "9. How to Manage After Going Live", level=1)

add_body(doc, (
    "Once your CGA platform is deployed and accessible to users, ongoing management becomes essential. "
    "The admin panel at /admin serves as your primary control center for day-to-day operations. This "
    "section covers all aspects of post-launch management."
))

add_heading(doc, "Managing Services", level=2)
add_body(doc, (
    "Navigate to Admin > Services in the admin panel. Here you can view all services, edit any service "
    "name, description, base price, feature list, and package tiers (Basic, Standard, Premium). You can "
    "also toggle service visibility to temporarily hide a service without deleting it. Changes made "
    "through the admin panel are reflected immediately on the public-facing website."
))

add_heading(doc, "Managing Orders", level=2)
add_body(doc, (
    "The Orders page (Admin > Orders) displays all customer orders with their current status. You can "
    "filter by status (pending, confirmed, in-progress, completed, cancelled), update order statuses, "
    "assign staff members to specific orders, add internal notes, and communicate delivery timelines. "
    "When an order status changes, the customer sees the update in their dashboard."
))

add_heading(doc, "Managing Users", level=2)
add_body(doc, (
    "The Users page (Admin > Users) shows all registered accounts. You can view user details, change "
    "user roles (between user and admin), activate or deactivate accounts, and view order history for "
    "each user. This is useful for onboarding team members or handling account issues."
))

add_heading(doc, "Database Management", level=2)
add_body(doc, (
    "For direct database access, use the Supabase Dashboard. Navigate to the Table Editor to view, "
    "add, edit, or delete records in any table. The SQL Editor allows you to run custom queries for "
    "bulk operations, data exports, or schema modifications. Always back up your data before making "
    "significant changes."
))

add_heading(doc, "Payment Management", level=2)
add_body(doc, (
    "Payment transactions are processed through SSLCOMMERZ. Log in to your SSLCOMMERZ merchant dashboard "
    "at sslcommerz.com to view transaction details, issue refunds, and download financial reports. "
    "Ensure your SSLCOMMERZ sandbox credentials are replaced with live production credentials before "
    "going live."
))

add_heading(doc, "Monitoring and Analytics", level=2)
add_body(doc, (
    "Vercel provides built-in Analytics, Logs, and Speed Insights for free. The Analytics dashboard "
    "shows page views, visitor demographics, and traffic sources. The Logs section displays real-time "
    "server logs and error traces. Speed Insights measures Core Web Vitals and page load performance."
))

add_heading(doc, "Code Updates", level=2)
add_body(doc, (
    "To update the live site, simply make your changes locally, commit to Git, and push to the main "
    "branch. Vercel automatically detects the push and triggers a new deployment. The entire process "
    "typically completes within two to three minutes. If a deployment fails, check the Vercel build "
    "logs for error details."
))

# ── Section 10: What Real Data You Need to Provide ─────────────────────

add_heading(doc, "10. What Real Data You Need to Provide", level=1)

add_body(doc, (
    "Before going live, certain real data must be provided to replace placeholder content. The table "
    "below serves as a checklist of all items you need to prepare."
), indent=False)

checklist_headers = ["Item", "Status", "Where to Use"]
checklist_rows = [
    ["Logo (PNG)", "Not provided yet", "Navbar, sidebar, login, favicon"],
    ["Brand colors", "Using default red-500", "Tailwind config / globals.css"],
    ["Real SSLCOMMERZ credentials", "Using mock/placeholder", ".env (SSLCOMMERZ_STORE_ID, SSLCOMMERZ_STORE_PASSWD)"],
    ["Real service prices (BDT)", "Using placeholder prices", "Admin Panel > Services OR seed script"],
    ["Real delivery times", "Using placeholder hours", "Admin Panel > Services > Edit packages"],
    ["YouTube portfolio links", "Not provided", "Services page, homepage"],
    ["Testimonials / reviews", "Not provided", "Homepage, services page"],
    ["FAQ content", "Not provided", "Can add to services pages"],
    ["WhatsApp number", "Not provided", "Footer, WhatsApp floating button"],
    ["Social media links", "Not provided", "Footer"],
    ["Custom domain", "Not provided", "Vercel settings"],
    ["Real coupon codes", "Not provided", "Database (Coupon table)"],
]
add_table(doc, checklist_headers, checklist_rows, col_widths=[2.2, 2.0, 2.8])
add_empty_line(doc)

# ── Section 11: 14-Day Development Roadmap ─────────────────────────────

add_heading(doc, "11. 14-Day Development Roadmap", level=1)

add_body(doc, (
    "The project follows a structured 14-day development plan. The table below shows the current "
    "status of each phase."
), indent=False)

roadmap_headers = ["Day", "Feature", "Status"]
roadmap_rows = [
    ["Day 1\u20133", "Setup, Auth, Services, UI", "Complete"],
    ["Day 4", "Order Form + Payment", "Complete"],
    ["Day 5", "Admin Panel", "Complete"],
    ["Day 6", "Blog / Content System", "Starting"],
    ["Day 7", "Chat Support System", "Pending"],
    ["Day 8", "AI Content Widget", "Pending"],
    ["Day 9", "Referral System Enhancement", "Pending"],
    ["Day 10", "Announcements & Notifications", "Pending"],
    ["Day 11", "PWA + Performance", "Pending"],
    ["Day 12", "i18n (Bengali / English)", "Pending"],
    ["Day 13", "Testing & Bug Fixes", "Pending"],
    ["Day 14", "Final Deployment", "Pending"],
]
add_table(doc, roadmap_headers, roadmap_rows, col_widths=[1.5, 3.5, 2.0])
add_empty_line(doc)

# ── Section 12: Important File Locations ───────────────────────────────

add_heading(doc, "12. Important File Locations (Quick Reference)", level=1)

add_body(doc, (
    "Use this table as a quick reference for finding the right file when you need to make changes "
    "to any part of the application."
), indent=False)

files_headers = ["What to Change", "File Path"]
files_rows = [
    ["Site title & meta", "src/app/layout.tsx"],
    ["Homepage content", "src/app/page.tsx"],
    ["Navigation bar", "src/components/layout/navbar.tsx"],
    ["Footer", "src/components/layout/footer.tsx"],
    ["Service pages", "src/app/services/[slug]/page.tsx"],
    ["Pricing page", "src/app/pricing/page.tsx"],
    ["Dashboard sidebar", "src/components/dashboard/sidebar.tsx"],
    ["Admin sidebar", "src/components/admin/admin-sidebar.tsx"],
    ["Order form", "src/components/orders/order-form.tsx"],
    ["API routes", "src/app/api/*/route.ts"],
    ["Database schema", "prisma/schema.prisma"],
    ["Auth logic", "src/lib/auth.ts, src/app/api/auth/*/route.ts"],
    ["Styles & colors", "src/app/globals.css, tailwind.config.ts"],
    ["Environment vars", ".env (never commit)"],
    ["Seed data", "scripts/seed-services.ts"],
]
add_table(doc, files_headers, files_rows, col_widths=[2.5, 4.5])

# ── Save ─────────────────────────────────────────────────────────────────

output_path = "/home/z/my-project/download/CGA-Project-Guide.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")